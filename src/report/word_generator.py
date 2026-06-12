"""
报告生成 — Word 整改建议报告
使用 python-docx 生成结构化 Word 整改建议报告。
"""

import io
import logging
from datetime import datetime, timezone
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger(__name__)


class WordReportGenerator:
    """Word 整改建议报告生成器"""

    def generate(
        self,
        run_id: str,
        timestamp: str,
        rule_set_version: str,
        total_items: int,
        hit_items: list[dict],
        risk_distribution: dict[str, dict],
        top_risk_categories: list[dict],
        output_path: str | None = None,
    ) -> bytes:
        """
        生成 Word 报告。

        Args:
            run_id: 演练记录ID
            timestamp: 执行时间（ISO 8601）
            rule_set_version: 规则集版本
            total_items: 总费用条数
            hit_items: 命中费用明细列表
            risk_distribution: 风险分布字典
            top_risk_categories: TOP 高风险类别列表
            output_path: 若指定则写入文件；否则返回 bytes

        Returns:
            DOCX 字节流
        """
        doc = Document()

        # ---- 全局样式（Normal 风格）----
        try:
            style = doc.styles["Normal"]
            style.font.name = "SimSun"
            style.font.size = Pt(12)
        except Exception:
            pass  # 非关键，降级

        # ---- 封面标题 ----
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("医保飞检规则演练整改建议报告")
        run.bold = True
        run.font.size = 36  # 18pt
        doc.add_paragraph()

        # ---- 元信息 ----
        meta = [
            ("演练记录ID", run_id),
            ("执行时间", timestamp),
            ("规则集版本", rule_set_version),
            ("费用总条数", str(total_items)),
            ("高风险项数量", str(len(hit_items))),
        ]
        for k, v in meta:
            p = doc.add_paragraph()
            p.add_run(f"{k}：").bold = True
            p.add_run(v)

        doc.add_paragraph()

        # ---- 一、高风险项清单 ----
        doc.add_heading("一、高风险项清单", level=1)

        if not hit_items:
            doc.add_paragraph("本次演练未发现高风险项。")
        else:
            # 构建表格
            headers = ["费用ID", "类别", "金额", "风险分", "规则名称", "匹配条件"]
            table_data = [headers]
            for item in hit_items:
                table_data.append([
                    str(item.get("item_id", "")),
                    str(item.get("category", "")),
                    f"¥{item.get('amount', 0):.2f}",
                    f"{item.get('risk_score', 0):.1f}",
                    str(item.get("name", "")),
                    str(item.get("matched_condition", "")),
                ])

            tbl = doc.add_table(rows=len(table_data), cols=len(headers))
            tbl.style = "Table Grid"

            for ri, row in enumerate(table_data):
                for ci, cell_text in enumerate(row):
                    cell = tbl.rows[ri].cells[ci]
                    cell.text = cell_text
                    # 表头加粗
                    if ri == 0:
                        cell.paragraphs[0].runs[0].bold = True

            doc.add_paragraph()

        # ---- 二、TOP 高风险类别 ----
        doc.add_heading("二、TOP 高风险类别", level=1)

        if not top_risk_categories:
            doc.add_paragraph("无高风险类别数据。")
        else:
            cat_headers = ["类别", "命中次数", "平均风险分", "风险金额合计"]
            cat_data = [cat_headers]
            for cat in top_risk_categories:
                cat_data.append([
                    str(cat.get("category", "")),
                    str(cat.get("hit_count", 0)),
                    f"{cat.get('avg_score', 0):.1f}",
                    f"¥{cat.get('total_amount', 0):.2f}",
                ])

            cat_tbl = doc.add_table(rows=len(cat_data), cols=len(cat_headers))
            cat_tbl.style = "Table Grid"

            for ri, row in enumerate(cat_data):
                for ci, cell_text in enumerate(row):
                    cell = cat_tbl.rows[ri].cells[ci]
                    cell.text = cell_text
                    if ri == 0:
                        cell.paragraphs[0].runs[0].bold = True

            doc.add_paragraph()

        # ---- 三、整改建议 ----
        doc.add_heading("三、整改建议", level=1)

        suggestions = self._generate_suggestions(top_risk_categories, risk_distribution)
        for sugg in suggestions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(sugg)

        # ---- 页脚 ----
        doc.add_paragraph()
        footer_p = doc.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer_p.add_run(
            f"生成时间：{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}　　"
            f"演练ID：{run_id}"
        ).font.size = 18  # 9pt

        # 保存
        if output_path:
            doc.save(output_path)
            logger.info("Word 报告已生成: %s", output_path)

        # 返回字节
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _generate_suggestions(
        self,
        top_risk_categories: list[dict],
        risk_distribution: dict,
    ) -> list[str]:
        """根据风险分布生成整改建议"""
        suggestions = []

        for cat in (top_risk_categories or []):
            cat_name = cat.get("category", "未知类别")
            avg_score = cat.get("avg_score", 0)
            hit_count = cat.get("hit_count", 0)

            if avg_score >= 80:
                level = "高危"
            elif avg_score >= 60:
                level = "中高危"
            elif avg_score >= 40:
                level = "中危"
            else:
                level = "低危"

            sugg = (
                f"【{cat_name}】风险等级：{level}，命中 {hit_count} 条，"
                f"平均风险分 {avg_score:.1f}。建议立即组织科室会诊，"
                f"核查费用明细是否合规，及时提交整改说明。"
            )
            suggestions.append(sugg)

        if not suggestions:
            suggestions.append("本次演练未发现明显违规项，建议持续监控费用变化趋势。")

        return suggestions
